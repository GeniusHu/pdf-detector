#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document processing service for extracting and filtering content
Supports PDF and Word (.docx) documents with page range configuration
"""

import asyncio
import os
import time
from typing import List, Tuple, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path
import logging

# Import our document processor
import sys
# Add project root to path
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from document_processor import DocumentProcessor, DocumentContent, Paragraph
from enhanced_pdf_extractor import EnhancedPDFTextExtractor, TextExtractionConfig
from models.api_models import ContentFilter, FileStatistics

logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Document content container - compatible with web API"""
    file_path: str
    file_type: str                           # pdf, docx
    paragraphs: List[Paragraph]             # List of processed paragraphs
    lines: List[Tuple[str, int, int]]       # Original lines (text, page, line_number)
    stats: FileStatistics
    extraction_time: float


class DocumentService:
    """Document processing service with support for PDF and Word"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def _parse_page_range(self, page_range_str: Optional[str]) -> Optional[Tuple[int, int]]:
        """
        Parse page range string like "1-146" to tuple (1, 146)

        Args:
            page_range_str: Page range string or None

        Returns:
            Tuple[int, int] or None
        """
        if not page_range_str:
            return None

        try:
            parts = page_range_str.strip().split('-')
            if len(parts) == 2:
                start = int(parts[0].strip())
                end = int(parts[1].strip())
                if start > 0 and end >= start:
                    return (start, end)
        except (ValueError, AttributeError):
            pass

        logger.warning(f"Invalid page range format: {page_range_str}")
        return None

    async def extract_document_content(
        self,
        file_path: str,
        content_filter: ContentFilter = ContentFilter.MAIN_CONTENT_ONLY,
        page_range: Optional[str] = None,
        task_id: Optional[str] = None,
        progress_callback: Optional[callable] = None
    ) -> DocumentContent:
        """
        Extract content from document with configurable filtering

        Args:
            file_path: Path to document file (PDF or Word)
            content_filter: Content filtering option
            page_range: Optional page range string (e.g., "1-146")
            task_id: Optional task ID for progress tracking
            progress_callback: Optional progress callback function

        Returns:
            DocumentContent: Extracted content with statistics
        """
        self.logger.info(f"[DocumentService] Starting extraction: {file_path}")
        start_time = time.time()

        try:
            # Parse page range
            parsed_page_range = self._parse_page_range(page_range)

            # Create extraction configuration based on filter option
            config = self._create_extraction_config(content_filter, parsed_page_range)

            # Use the new DocumentProcessor
            processor = DocumentProcessor(config)

            # Process document - 使用 asyncio.to_thread 在后台线程执行同步操作，避免阻塞事件循环
            doc_content = await asyncio.to_thread(processor.process, file_path)

            # Also get original lines for backward compatibility - 同样在后台线程执行
            lines = await asyncio.to_thread(self._extract_lines, file_path, config)

            # Calculate file size
            file_size = os.path.getsize(file_path)
            file_size_mb = round(file_size / (1024 * 1024), 2)

            # Detect page count
            total_pages = self._get_page_count(file_path)

            # Create FileStatistics
            file_stats = FileStatistics(
                file_path=file_path,
                file_size_mb=file_size_mb,
                total_pages=total_pages,
                total_lines=len(lines),
                main_content_lines=len(lines),  # After filtering
                filtered_lines=0,  # Will be calculated
                total_chars=doc_content.total_raw_chars,
                processing_time_seconds=time.time() - start_time
            )

            # Create content object
            content = DocumentContent(
                file_path=file_path,
                file_type=doc_content.file_type,
                paragraphs=doc_content.paragraphs,
                lines=lines,
                stats=file_stats,
                extraction_time=time.time() - start_time
            )

            self.logger.info(f"[DocumentService] Extraction completed: "
                           f"{content.file_type.upper()} - {len(lines)} lines, "
                           f"{len(doc_content.paragraphs)} paragraphs, "
                           f"{doc_content.total_clean_chars} clean chars")

            return content

        except Exception as e:
            self.logger.error(f"Error extracting document content from {file_path}: {str(e)}", exc_info=True)
            raise

    def _extract_lines(self, file_path: str, config) -> List[Tuple[str, int, int]]:
        """Extract lines from document using the appropriate extractor"""
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.pdf':
            extractor = EnhancedPDFTextExtractor(config)
            return extractor.extract_main_text_lines(file_path)
        elif file_ext == '.docx':
            from word_extractor import WordExtractor
            extractor = WordExtractor(config)
            return extractor.extract_text_with_positions(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def _get_page_count(self, file_path: str) -> int:
        """Get page/paragraph count of document"""
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return len(pdf.pages)
            except:
                return 0
        elif file_ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                return len(doc.paragraphs)
            except:
                return 0
        return 0

    def _create_extraction_config(
        self,
        content_filter: ContentFilter,
        page_range: Optional[Tuple[int, int]] = None
    ) -> TextExtractionConfig:
        """Create extraction configuration based on filter option"""

        # Base configuration - main content only by default
        config = TextExtractionConfig(
            include_references=False,
            include_footnotes=False,
            include_citations=False,
            include_page_numbers=False,
            include_headers_footers=False,
            include_annotations=False,
            min_line_length=5,
            remove_duplicate_lines=True,
            page_range=page_range
        )

        # Adjust based on content filter
        if content_filter == ContentFilter.ALL_CONTENT:
            config.include_references = True
            config.include_footnotes = True
            config.include_citations = True
            config.include_page_numbers = True
            config.include_headers_footers = True
            config.min_line_length = 5

        elif content_filter == ContentFilter.INCLUDE_REFERENCES:
            config.include_references = True
            config.min_line_length = 10

        elif content_filter == ContentFilter.INCLUDE_CITATIONS:
            config.include_citations = True
            config.min_line_length = 10

        # MAIN_CONTENT_ONLY uses the default config above

        return config

    async def validate_document_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validate document file and return basic information

        Args:
            file_path: Path to document file

        Returns:
            Dict: Validation results
        """
        try:
            if not os.path.exists(file_path):
                return {"valid": False, "error": "File not found"}

            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in ['.pdf', '.docx']:
                return {"valid": False, "error": f"Unsupported file type: {file_ext}"}

            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return {"valid": False, "error": "Empty file"}

            # Get page/paragraph count
            page_count = self._get_page_count(file_path)

            return {
                "valid": True,
                "file_path": file_path,
                "file_type": "pdf" if file_ext == ".pdf" else "docx",
                "file_size_mb": round(file_size / (1024 * 1024), 2),
                "page_count": page_count
            }

        except Exception as e:
            return {
                "valid": False,
                "error": f"Invalid document file: {str(e)}"
            }

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return [".pdf", ".docx"]

    def estimate_processing_time(
        self,
        file_size_mb: float,
        page_count: int,
        content_filter: ContentFilter,
        processing_mode: str
    ) -> float:
        """
        Estimate processing time based on file characteristics

        Args:
            file_size_mb: File size in MB
            page_count: Number of pages
            content_filter: Content filtering option
            processing_mode: Processing mode

        Returns:
            float: Estimated processing time in seconds
        """
        # Base processing time (seconds per page)
        base_time_per_page = 0.5

        # Adjust based on content filter
        if content_filter == ContentFilter.ALL_CONTENT:
            filter_multiplier = 1.5
        elif content_filter == ContentFilter.MAIN_CONTENT_ONLY:
            filter_multiplier = 1.0
        else:
            filter_multiplier = 1.2

        # Adjust based on processing mode
        if processing_mode == "ultra_fast":
            mode_multiplier = 0.3
        elif processing_mode == "fast":
            mode_multiplier = 0.6
        else:
            mode_multiplier = 1.0

        # Calculate estimated time
        estimated_time = (
            page_count * base_time_per_page *
            filter_multiplier * mode_multiplier
        )

        # Add overhead for large files
        if file_size_mb > 50:
            estimated_time *= 1.5
        elif file_size_mb > 20:
            estimated_time *= 1.2

        return max(estimated_time, 5.0)  # Minimum 5 seconds
