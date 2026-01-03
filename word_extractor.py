#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档提取器
支持.docx格式的Word文档文本提取
"""

import re
from typing import List, Tuple, Optional
from dataclasses import dataclass

try:
    from docx import Document
except ImportError:
    raise ImportError("需要安装 python-docx 库。请运行: pip install python-docx")

from document_extractor import BaseDocumentExtractor


@dataclass
class WordExtractionConfig:
    """Word文档提取配置"""
    include_tables: bool = True           # 是否包含表格内容
    include_headers_footers: bool = False  # 是否包含页眉页脚
    include_footnotes_endnotes: bool = False  # 是否包含脚注和尾注
    min_paragraph_length: int = 5         # 最小段落长度
    remove_duplicate_lines: bool = True   # 是否去除重复行
    page_range: Tuple[int, int] = None    # 段落范围 (start, end)


class WordExtractor(BaseDocumentExtractor):
    """Word文档提取器"""

    def __init__(self, config=None):
        """
        初始化Word提取器

        Args:
            config: 提取配置 (TextExtractionConfig 或 WordExtractionConfig)
        """
        self.config = config or WordExtractionConfig()

    def _get_config_attr(self, attr, default=None):
        """获取配置属性，兼容不同的配置类型"""
        if hasattr(self.config, attr):
            return getattr(self.config, attr)

        # 属性映射：TextExtractionConfig -> WordExtractor 兼容
        attr_mapping = {
            'min_line_length': 'min_paragraph_length',
            'min_paragraph_length': 'min_line_length',
            'remove_duplicate_lines': 'remove_duplicate_lines',
            'include_references': 'include_references',
            'include_footnotes': 'include_footnotes_endnotes',
            'include_headers_footers': 'include_headers_footers',
            'page_range': 'page_range',
        }

        # 尝试映射后的属性名
        for old_attr, new_attr in attr_mapping.items():
            if attr == old_attr and hasattr(self.config, new_attr):
                return getattr(self.config, new_attr)
            if attr == new_attr and hasattr(self.config, old_attr):
                return getattr(self.config, old_attr)

        return default

    def supports_file_type(self, file_path: str) -> bool:
        """判断是否支持该文件类型"""
        return file_path.lower().endswith('.docx')

    def extract_text_with_positions(self, file_path: str) -> List[Tuple[str, int, int]]:
        """
        提取Word文档文本及其位置信息

        Args:
            file_path: Word文档路径

        Returns:
            List[Tuple[str, int, int]]: (文本, 段落号, 句子号) 的列表
        """
        print(f"\n{'='*60}")
        print(f"[WORD] ===== STARTING EXTRACTION: {file_path} =====")
        print(f"{'='*60}")

        doc = Document(file_path)

        text_lines = []
        paragraph_count = 0
        skipped_count = 0

        # 应用段落范围限制
        para_start, para_end = 1, len(doc.paragraphs)
        page_range = self._get_config_attr('page_range', None)
        if page_range:
            para_start, para_end = page_range
            para_end = min(para_end, len(doc.paragraphs))
            print(f"[WORD] Paragraph range: {para_start}-{para_end} (total: {len(doc.paragraphs)} paragraphs)")
        else:
            print(f"[WORD] Total paragraphs: {len(doc.paragraphs)}")

        # 获取最小段落长度（兼容不同配置类型）
        min_length = self._get_config_attr('min_paragraph_length') or self._get_config_attr('min_line_length', 5)

        # 提取正文段落
        for idx, paragraph in enumerate(doc.paragraphs, 1):
            # 跳过范围外的段落
            if idx < para_start or idx > para_end:
                continue

            text = paragraph.text.strip()

            # 跳过空段落
            if not text:
                continue

            # 跳过太短的段落
            if len(text) < min_length:
                skipped_count += 1
                continue

            # 检查是否是脚注/参考文献
            if self._is_footnote_line(text):
                skipped_count += 1
                continue

            # 标准化文本
            normalized_text = self._normalize_text(text)

            if normalized_text:
                # Word没有真实分页，用段落号模拟页码
                text_lines.append((normalized_text, idx, 1))
                paragraph_count += 1

            # 每100个段落报告进度
            if idx % 100 == 0 or idx == para_end:
                print(f"[WORD] Processed {idx}/{para_end} paragraphs, {paragraph_count} lines extracted")

        # 提取表格内容（如果启用）
        include_tables = self._get_config_attr('include_tables', True)
        if include_tables:
            table_count = 0
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if text and len(text) >= min_length:
                            normalized_text = self._normalize_text(text)
                            if normalized_text:
                                # 表格内容页码设为9999表示是表格
                                text_lines.append((normalized_text, 9999, table_idx))
                                table_count += 1

            print(f"[WORD] Extracted {table_count} lines from tables")

        # 去除重复行
        remove_dup = self._get_config_attr('remove_duplicate_lines', True)
        if remove_dup:
            text_lines = self._remove_duplicates(text_lines)

        print(f"[WORD] COMPLETED: {paragraph_count} lines extracted, {skipped_count} lines skipped")
        print(f"{'='*60}\n")

        return text_lines

    def _is_footnote_line(self, text: str) -> bool:
        """判断是否为脚注/参考文献行"""
        # 中文脚注模式
        footnote_patterns = [
            r'参见.*第\d+页',
            r'详见.*第\d+页',
            r'出版社.*年版第\d+页',
            r'人民出版社.*年版',
            r'中央文献出版社.*年版',
            r'文献出版社.*年版',
            r'学习出版社.*年版',
            r'第\d+卷.*第\d+页',
            r'\d{4}年.*版',
            r'年版.*第\d+页',
            r'ISBN',
            r'ISSN',
            r'DOI:',
        ]

        for pattern in footnote_patterns:
            if re.search(pattern, text):
                return True

        return False

    def _normalize_text(self, text: str) -> str:
        """标准化文本"""
        # 去除多余空格
        text = re.sub(r'\s+', ' ', text)
        # 去除首尾空格
        text = text.strip()
        return text

    def _remove_duplicates(self, lines: List[Tuple[str, int, int]]) -> List[Tuple[str, int, int]]:
        """去除重复行"""
        seen_texts = set()
        unique_lines = []

        for text, page, line in lines:
            if text not in seen_texts:
                seen_texts.add(text)
                unique_lines.append((text, page, line))

        return unique_lines
